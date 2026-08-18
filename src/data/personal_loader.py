from pathlib import Path
from docx import Document
from datetime import datetime


class PersonalLoader:

    def __init__(
        self,
        output_dir="data/personal_docs"
    ):
        self.output_dir = Path(output_dir)

        self.output_dir.mkdir(
            parents=True,
            exist_ok=True
        )


    # ==========================
    # Text cleaning
    # ==========================

    def clean_text(self, text):

        if not text:
            return ""

        text = text.replace("\n", " ")

        text = " ".join(
            text.split()
        )

        return text.strip()



    # ==========================
    # Detect document type
    # ==========================

    def detect_document_type(
        self,
        full_text
    ):

        text = full_text.upper()


        if (
            "HỢP ĐỒNG" in text
            and
            (
                "CHUYỂN NHƯỢNG" in text
                or
                "ĐẶT CỌC" in text
            )
        ):

            return {

                "document_type":
                    "contract",

                "category":
                    "real_estate",

                "domain":
                    "legal"

            }



        if "HỢP ĐỒNG" in text:

            return {

                "document_type":
                    "contract",

                "category":
                    "general",

                "domain":
                    "legal"

            }



        return {

            "document_type":
                "document",

            "category":
                "general",

            "domain":
                "unknown"

        }



    # ==========================
    # Clean merged table row
    # ==========================

    def clean_table_row(
        self,
        row
    ):

        values = []


        for cell in row.cells:

            text = self.clean_text(
                cell.text
            )


            if not text:
                continue


            # loại duplicate do merge cell

            if text not in values:

                values.append(text)



        return values



    # ==========================
    # Convert table to markdown
    # ==========================

    def table_to_markdown(
        self,
        table
    ):

        rows = []


        for row in table.rows:

            values = self.clean_table_row(
                row
            )


            if values:

                rows.append(values)



        if not rows:

            return ""



        # Nếu toàn bộ row chỉ có 1 cell

        max_columns = max(
            len(row)
            for row in rows
        )


        if max_columns == 1:

            markdown = [
                "| Nội dung |",
                "|---|"
            ]


            for row in rows:

                markdown.append(
                    f"| {row[0]} |"
                )


            return "\n".join(markdown)



        markdown = []



        header = rows[0]


        markdown.append(
            "| "
            +
            " | ".join(header)
            +
            " |"
        )


        markdown.append(
            "| "
            +
            " | ".join(
                ["---"] * len(header)
            )
            +
            " |"
        )



        for row in rows[1:]:


            while len(row) < len(header):

                row.append("")



            markdown.append(
                "| "
                +
                " | ".join(row)
                +
                " |"
            )



        return "\n".join(markdown)



    # ==========================
    # DOCX -> Markdown
    # ==========================

    def convert_docx_to_markdown(
        self,
        input_file
    ):

        path = Path(input_file)


        if not path.exists():

            raise FileNotFoundError(
                f"Không tìm thấy file: {input_file}"
            )


        if path.suffix.lower() != ".docx":

            raise ValueError(
                "Chỉ hỗ trợ DOCX"
            )



        doc = Document(
            path
        )



        # lấy text toàn bộ để detect

        all_text = []


        for para in doc.paragraphs:

            text = self.clean_text(
                para.text
            )

            if text:

                all_text.append(text)



        document_info = self.detect_document_type(
            "\n".join(all_text)
        )



        markdown = []


        # Metadata

        markdown.append(
            f"<!-- DOCUMENT_TYPE: {document_info['document_type']} -->"
        )

        markdown.append(
            f"<!-- CATEGORY: {document_info['category']} -->"
        )

        markdown.append(
            f"<!-- DOMAIN: {document_info['domain']} -->"
        )

        markdown.append(
            f"<!-- SOURCE_FILE: {path.name} -->"
        )

        markdown.append(
            f"<!-- CONVERTED: {datetime.now()} -->"
        )

        markdown.append("")



        # Paragraph

        for para in doc.paragraphs:


            text = self.clean_text(
                para.text
            )


            if not text:

                continue



            upper = text.upper()



            if (
                upper.startswith("ĐIỀU ")
                or
                "BIÊN NHẬN THANH TOÁN"
                in upper
            ):

                markdown.append("")

                markdown.append(
                    f"## {text}"
                )

                markdown.append("")


            else:

                markdown.append(
                    text
                )

                markdown.append("")



        # Tables

        for index, table in enumerate(
            doc.tables
        ):


            markdown.append("")


            markdown.append(
                "<!-- TABLE_TYPE: financial -->"
            )


            table_md = self.table_to_markdown(
                table
            )


            if table_md:

                markdown.append(
                    table_md
                )


            markdown.append("")



        output = (
            self.output_dir
            /
            f"{path.stem}.md"
        )



        with open(
            output,
            "w",
            encoding="utf-8"
        ) as f:

            f.write(
                "\n".join(markdown)
            )



        return output




if __name__ == "__main__":


    loader = PersonalLoader()


    output = loader.convert_docx_to_markdown(
        "data/personal_docs/A5.16.03.docx"
    )


    print(
        "Created:"
    )

    print(output)